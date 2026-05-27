"""
FastAPI 合同管理 API 路由
"""
from fastapi import APIRouter, Depends, File, UploadFile, HTTPException, status, Query
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from database import get_db
from schemas import ContractResponse, ContractUpload
from models import Contract, AsyncTask, User
from services.business_logic import invoice_service
from utils.file_handler import save_upload_file, get_file_extension, get_file_size
from tasks.celery_tasks import contract_analyze_risks
from config import settings
from deps import get_current_user
from utils.audit import write_audit_log
import os
import uuid

router = APIRouter(prefix="/api/contracts", tags=["contracts"])

CONTRACT_EXTENSIONS = {".pdf", ".docx"}


def extract_pdf_text(file_path: str) -> str:
    """从 PDF 提取文本"""
    try:
        import PyPDF2
        text = ""
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
        return text
    except:
        return ""


def extract_docx_text(file_path: str) -> str:
    """从 DOCX 提取文本"""
    try:
        from docx import Document

        document = Document(file_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        table_cells.append(cell.text.strip())
        return "\n".join(paragraphs + table_cells)
    except Exception:
        return ""


def extract_contract_text(file_path: str) -> str:
    """从合同文件中提取文本"""
    extension = get_file_extension(file_path)
    if extension == ".pdf":
        return extract_pdf_text(file_path)
    if extension == ".docx":
        return extract_docx_text(file_path)
    return ""


@router.post("/upload")
async def upload_contract(
    file: UploadFile = File(...),
    contract_name: str = None,
    supplier_name: str = None,
    amount: float = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """上传合同文件"""
    # 检查文件类型
    if get_file_extension(file.filename) not in CONTRACT_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only PDF and DOCX contract files are supported"
        )
    
    # 保存文件
    file_path = save_upload_file(file, subfolder="contracts")
    file_size = get_file_size(file_path)
    
    # 检查文件大小
    if file_size > settings.MAX_FILE_SIZE:
        os.remove(file_path)
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File size exceeds 50MB limit"
        )
    
    # 创建合同记录
    contract = Contract(
        contract_number=f"CTR-{uuid.uuid4().hex[:10].upper()}",
        contract_name=contract_name or file.filename,
        supplier_name=supplier_name,
        amount=amount,
        file_path=file_path,
        file_name=file.filename,
        file_size=file_size,
        upload_user_id=current_user.id,
        status="pending"
    )
    
    db.add(contract)
    db.commit()
    db.refresh(contract)
    
    # 提取合同文本
    contract_text = extract_contract_text(file_path)
    if not contract_text.strip():
        contract.status = "error"
        contract.analysis_error = "未能从合同文件中提取可分析文本，请上传可复制文本的 PDF 或 DOCX 文件"
        db.commit()
        write_audit_log(
            db,
            action="contract_upload_failed",
            resource_type="contract",
            resource_id=contract.id,
            user_id=current_user.id,
            changes={
                "contract_number": contract.contract_number,
                "contract_name": contract.contract_name,
                "file_name": file.filename,
                "error_message": contract.analysis_error,
            },
        )
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=contract.analysis_error
        )
    
    task_id = str(uuid.uuid4())
    db.add(AsyncTask(
        task_id=task_id,
        task_type="contract_analysis",
        status="pending",
        resource_type="contract",
        resource_id=contract.id,
    ))
    db.commit()

    write_audit_log(
        db,
        action="contract_uploaded",
        resource_type="contract",
        resource_id=contract.id,
        user_id=current_user.id,
        changes={
            "contract_number": contract.contract_number,
            "contract_name": contract.contract_name,
            "file_name": file.filename,
        },
    )

    if settings.BACKGROUND_TASK_MODE == "inline":
        contract_analyze_risks.apply(args=[contract.id, contract_text[:5000]], task_id=task_id)
    else:
        contract_analyze_risks.apply_async(args=[contract.id, contract_text[:5000]], task_id=task_id)
    
    return {
        "contract_id": contract.id,
        "contract_number": contract.contract_number,
        "file_name": file.filename,
        "task_id": task_id,
        "status": "uploaded",
        "message": "Contract uploaded successfully, analysis started"
    }


@router.get("/{contract_id}", response_model=ContractResponse)
def get_contract(
    contract_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取合同详情"""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    
    if not contract or (current_user.role != "admin" and contract.upload_user_id != current_user.id):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Contract not found"
        )
    
    return contract


@router.get("/{contract_id}/file")
def download_contract_file(
    contract_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """下载合同原始文件"""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()

    if not contract or (current_user.role != "admin" and contract.upload_user_id != current_user.id):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Contract not found"
        )

    if not contract.file_path or not os.path.exists(contract.file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Contract file not found"
        )

    return FileResponse(
        contract.file_path,
        filename=contract.file_name or f"contract-{contract_id}{get_file_extension(contract.file_path)}",
        media_type="application/octet-stream",
    )


@router.get("/", response_model=list[ContractResponse])
def list_contracts(
    skip: int = Query(0, ge=0),
    limit: int = Query(10, ge=1, le=100),
    status: str = Query(None),
    risk_level: str = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """列表查询合同"""
    query = db.query(Contract)
    if current_user.role != "admin":
        query = query.filter(Contract.upload_user_id == current_user.id)
    
    if status:
        query = query.filter(Contract.status == status)
    
    if risk_level:
        query = query.filter(Contract.risk_level == risk_level)
    
    contracts = query.order_by(Contract.created_at.desc()).offset(skip).limit(limit).all()
    return contracts


@router.get("/{contract_id}/risks")
def get_contract_risks(
    contract_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取合同风险"""
    from models import ContractRisk
    
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    
    if not contract or (current_user.role != "admin" and contract.upload_user_id != current_user.id):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Contract not found"
        )
    
    risks = db.query(ContractRisk).filter(
        ContractRisk.contract_id == contract_id
    ).all()
    
    return {
        "contract_id": contract_id,
        "risk_score": contract.risk_score,
        "risk_level": contract.risk_level,
        "total_risks": len(risks),
        "risks": risks
    }


@router.get("/{contract_id}/analysis-status")
def get_analysis_status(
    contract_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """获取分析状态"""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    
    if not contract or (current_user.role != "admin" and contract.upload_user_id != current_user.id):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Contract not found"
        )
    
    return {
        "contract_id": contract_id,
        "status": contract.status,
        "risk_score": contract.risk_score,
        "risk_level": contract.risk_level,
        "analysis_started_at": contract.analysis_started_at,
        "analysis_completed_at": contract.analysis_completed_at,
        "error": contract.analysis_error
    }
